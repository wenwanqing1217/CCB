from docx import Document
import os

def extract_doc_content(docx_path, output_name):
    try:
        doc = Document(docx_path)
        print(f"\n{'='*60}")
        print(f"==== {output_name} ====")
        print(f"{'='*60}\n")
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"{i}: {para.text}")
        
        print(f"\n{'='*60}")
        print(f"==== {output_name} 表格 ====")
        print(f"{'='*60}\n")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_data))
                
    except Exception as e:
        print(f"错误读取 {docx_path}: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    extract_doc_content('撰写规范.docx', '撰写规范')
    print("\n\n")
    extract_doc_content('论文_完整版.docx', '论文完整版')
