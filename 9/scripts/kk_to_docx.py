#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将kk.md论文转换为Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_three_line_table(doc, headers, rows, col_widths=None):
    """添加三线表"""
    if not headers or not rows:
        return
    
    full_data = [headers] + rows
    table = doc.add_table(rows=len(full_data), cols=len(headers))
    
    # 设置表格样式
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(full_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text)
            
            # 设置字体
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置表头样式（上框线加粗）
    for cell in table.rows[0].cells:
        cell._tc.get_or_add_tcPr()
        
    # 设置表格宽度
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
    
    return table

def parse_markdown_to_docx(md_file, docx_file):
    """解析Markdown并转换为Word文档"""
    
    doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    table_lines = []
    in_table = False
    in_formula = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 数学公式
        if line.startswith('$$'):
            if in_formula:
                in_formula = False
                doc.add_paragraph()
            else:
                in_formula = True
            i += 1
            continue
        
        if in_formula:
            i += 1
            continue
        
        # 代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.line_spacing = 1.2
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 表格检测
        if '|' in line and line.count('|') >= 2:
            table_row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if table_row and not line.startswith('---'):
                if not in_table:
                    in_table = True
                    table_lines = []
                # 检查是否分隔行
                if not all(c == '-' or c == ':' for c in line.replace('|', '').replace(' ', '')):
                    table_lines.append(table_row)
            i += 1
            continue
        else:
            if in_table and table_lines:
                # 添加表格
                if len(table_lines) >= 2:
                    headers = table_lines[0]
                    rows = table_lines[1:]
                    add_three_line_table(doc, headers, rows)
                else:
                    # 普通段落式表格
                    pass
                table_lines = []
                in_table = False
        
        # 标题处理
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group(0))
            text = line.lstrip('# ').strip()
            
            if level == 1:
                # 一级标题居中
                h = doc.add_heading(text, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                # 二级标题左对齐
                h = doc.add_heading(text, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                h = doc.add_heading(text, level=level)
            i += 1
            continue
        
        # 分隔线
        if line.startswith('---') or line.startswith('==='):
            i += 1
            continue
        
        # 空行
        if not line.strip():
            i += 1
            continue
        
        # 普通段落
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        i += 1
    
    # 保存
    doc.save(docx_file)
    print(f"✓ 论文已生成: {docx_file}")

if __name__ == "__main__":
    parse_markdown_to_docx('kk.md', 'kk.docx')
