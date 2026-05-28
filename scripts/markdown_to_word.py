from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    if level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

def add_table(doc, table_data):
    """添加表格"""
    if not table_data:
        return
    
    rows = len(table_data)
    cols = len(table_data[0]) if rows > 0 else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = str(cell_text)

def parse_markdown_to_word(md_file_path, docx_file_path):
    """解析Markdown并写入Word文档"""
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按行处理
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    table_content = []
    
    for line in lines:
        line = line.rstrip()
        
        # 代码块开始/结束
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                add_code_block(doc, '\n'.join(code_content))
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # 标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group(0))
            text = line.lstrip('# ').strip()
            add_heading(doc, text, level=level if level <= 9 else 9)
            continue
        
        # 表格行
        if '|' in line and line.count('|') >= 2:
            # 简单的表格检测
            table_row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if table_row:
                table_content.append(table_row)
                continue
        
        # 如果之前有表格内容，现在处理
        if table_content:
            add_table(doc, table_content)
            table_content = []
        
        # 分隔线
        if line.startswith('---') or line.startswith('==='):
            doc.add_paragraph()
            continue
        
        # 普通段落
        if line.strip():
            # 检查粗体和斜体
            text = line
            # 简单处理，先不处理粗体斜体
            add_paragraph(doc, text)
        else:
            # 空行
            doc.add_paragraph()
    
    # 处理剩余的表格
    if table_content:
        add_table(doc, table_content)
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"Word文档已生成: {docx_file_path}")

if __name__ == "__main__":
    try:
        parse_markdown_to_word('论文_完整版.md', '6.docx')
        print("转换成功！")
    except Exception as e:
        print(f"转换失败: {e}")
        import traceback
        traceback.print_exc()
