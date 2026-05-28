from docx import Document
import sys
import os

doc_path = '2406910639郑茜-毕业论文 - 副本.docx'

try:
    doc = Document(doc_path)
    print("=== 毕业论文内容 ===\n")
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{para.text}\n")
    
    print("\n=== 表格内容 ===\n")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx + 1}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_data))
            
except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()
