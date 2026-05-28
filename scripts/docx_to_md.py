from docx import Document
import os

def docx_to_markdown(docx_path, output_path):
    try:
        doc = Document(docx_path)
        md_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text = para.text.strip()
                
                # 判断是否是标题
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading', '')
                    if level.isdigit():
                        md_content.append('#' * int(level) + ' ' + text + '\n')
                    else:
                        md_content.append('## ' + text + '\n')
                else:
                    md_content.append(text + '\n')
        
        # 处理表格
        for table_idx, table in enumerate(doc.tables):
            md_content.append(f'\n### 表格 {table_idx + 1}\n')
            for i, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    md_content.append('| ' + ' | '.join(row_data) + ' |\n')
                    md_content.append('| ' + ' | '.join(['---'] * len(row_data)) + ' |\n')
                else:
                    md_content.append('| ' + ' | '.join(row_data) + ' |\n')
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(md_content))
        
        print(f'转换成功！Markdown文件已保存到: {output_path}')
        return ''.join(md_content)
        
    except Exception as e:
        print(f'转换失败: {e}')
        import traceback
        traceback.print_exc()
        return None

if __name__ == '__main__':
    docx_path = '6.docx'
    output_path = '6.md'
    content = docx_to_markdown(docx_path, output_path)
